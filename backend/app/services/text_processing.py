"""Text extraction from uploaded reports — the front door Track A owns.

extract_text() turns a PDF/DOCX/TXT upload into plain text. The lab-value
flagging itself lives in the lab_triage skill (deterministic, cited tables);
this module only gets the text out of the file.
"""
from __future__ import annotations

from pathlib import Path

# Image suffixes go to the imaging/OpenAI-vision path, not text extraction.
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
TEXT_SUFFIXES = {".pdf", ".docx", ".txt", ".csv", ".json", ".md"}


def is_image(path: str) -> bool:
    return Path(path).suffix.lower() in IMAGE_SUFFIXES


def route_file(path: str) -> tuple[str, str]:
    """Decide how an upload reaches the extractor.

    Returns ("text", extracted_text) for anything with a usable text layer
    (digital PDF/DOCX/TXT — free, exact, no model), or ("image", path) for
    photos and SCANNED PDFs with no text layer (the vision model reads them;
    a PDF is rendered to page images downstream).
    """
    if is_image(path):
        return "image", path
    if Path(path).suffix.lower() == ".pdf":
        text = extract_text(path)
        if len(text.strip()) >= 40:        # has a real text layer
            return "text", text
        return "image", path               # scanned PDF -> vision
    return "text", extract_text(path)


def extract_text(path: str) -> str:
    """Pull text from a report. Empty string if there is no text layer
    (e.g. a scanned PDF) — the caller then routes to the vision extractor."""
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix == ".docx":
        return _from_docx(path)
    if suffix in {".txt", ".csv", ".json", ".md"}:
        return Path(path).read_text(errors="ignore")
    return ""


def _from_pdf(path: str) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return _from_pdf_pypdf(path)
    text = []
    with fitz.open(path) as doc:
        for page in doc:
            text.append(page.get_text())
    return "\n".join(text).strip()


def _from_pdf_pypdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    reader = PdfReader(path)
    return "\n".join((p.extract_text() or "") for p in reader.pages).strip()


def _from_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:                       # lab results are usually tables
        for row in table.rows:
            lines.append("\t".join(c.text for c in row.cells))
    return "\n".join(lines).strip()
